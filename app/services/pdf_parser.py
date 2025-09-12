import io, os, zipfile
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document

PARSER_VERSION = "sniff-2"   # <- shows up in /health and logs

class ResumeParser:
    """
    Filename-agnostic resume parser.
    - Detects PDF vs DOCX from bytes.
    - Falls back to extension only if sniffing is inconclusive.
    - NEVER requires the uploaded filename.
    """

    async def parse_resume(self, filename: Optional[str], file_bytes: bytes) -> str:
        if not file_bytes:
            raise ValueError("Empty file")
        print(f"[parser] version={PARSER_VERSION} filename={filename!r}")

        kind = self._detect_kind(file_bytes)  # primary: sniff bytes

        # secondary: extension hint (not required)
        if kind is None and filename:
            _, ext = os.path.splitext(filename.lower())
            if ext in {".pdf", ".docx"}:
                kind = ext[1:]

        if kind == "pdf":
            return self._parse_pdf(file_bytes)
        if kind == "docx":
            return self._parse_docx(file_bytes)

        raise ValueError("Unsupported or unknown file type.")

    # ---------- helpers ----------
    def _detect_kind(self, b: bytes) -> Optional[str]:
        # PDF starts with %PDF-
        if len(b) >= 5 and b[:5] == b"%PDF-":
            return "pdf"
        # DOCX is a zip with [Content_Types].xml and word/...
        if len(b) >= 2 and b[:2] == b"PK":
            try:
                with io.BytesIO(b) as bio, zipfile.ZipFile(bio) as zf:
                    names = set(zf.namelist())
                    if "[Content_Types].xml" in names and any(n.startswith("word/") for n in names):
                        return "docx"
            except zipfile.BadZipFile:
                pass
        return None

    def _parse_pdf(self, b: bytes) -> str:
        out = []
        with io.BytesIO(b) as bio:
            reader = PdfReader(bio)
            for page in reader.pages:
                try:
                    out.append(page.extract_text() or "")
                except Exception:
                    continue
        text = "\n".join(out).strip()
        if not text:
            raise ValueError("Unable to extract text from PDF")
        return text

    def _parse_docx(self, b: bytes) -> str:
        with io.BytesIO(b) as bio:
            doc = Document(bio)
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            parts.append(cell.text.strip())
        text = "\n".join(parts).strip()
        if not text:
            raise ValueError("Unable to extract text from DOCX")
        return text
